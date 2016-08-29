#! /usr/bin/env python3

##
## IMPORTS
##

import json
import re
from docx import Document
from bs4 import BeautifulSoup
from lxml import etree as et
from lxml import html as lhtml
from lxml.html.clean import Cleaner
import requests
import logging
import configparser

##
## CONSTANTS and VARIABLES
##

config = configparser.ConfigParser()
config.read("../data/default.ini")

if config["DEFAULT"]["log_level"] == "debug":
  logging.basicConfig(level=logging.DEBUG)
else:
  logging.basicConfig(level=logging.WARN)

HTML_IDS = config["scraping"]["html_ids"].split(",")
WRITE_DIR = config["DEFAULT"]["root_dir"] + config["scraping"]["write_dir"] + "/"
CLASSIFICATIONS = config["DEFAULT"]["classifications"].split(",") # ["unclassified", "sensitive but unclassified", "confidential", "secret", "noforn"]
PARA_MARKINGS = config["DEFAULT"]["para_markings"].split(",") # ["U", "SBU", "C", "S", "noforn"]

# doc loading -
country_doc = "../data/countries.json"
org_doc = "../data/organization.json"

with open(country_doc, "r") as d:
  c_dict = json.load(d)

c_list = list(c_dict.values())
logging.info("c_list loaded")

with open(org_doc, "r") as d:
  o_dict = json.load(d)

o_list = list(o_dict.values())
logging.info("o_list loaded")

##
## FUNCTIONS
##

def batchScrape(lst, html_ids):
  '''Accepts a list of URLs and an HTML element id'''
  for item in lst:
    d = Scrape2Docx()
    d.pipeline(item, html_ids)

def main():
  pass

##
## CLASSES
##

class Scrape2Docx(object):
  def __init__(self):
    self._b = None
    self._text = None
    self._text_list = None

  def fetch(self, url):
    self.url = url
    logging.info(url)
    r_ = requests.get(url)

    if r_.status_code != 200:
      logging.warn("returned status code: " + str(r_.status_code))
      return
    else:
      html = r_.text
      self._b = BeautifulSoup(html, 'html.parser')
      logging.info(url + " fetched...")

  def parse(self, html_ids):
    self.h_ids = html_ids
    if self._b is None:
      logging.warn("URL has not been fetched!") # to do
      return
    else:
      for id_ in self.h_ids:
        t_ = self._b.find(id=id_) # target
        if t_ is None:
          logging.warn(id_ + " not found in " + self.url)
        else:
          logging.info("Length of located content in scrape:" + str(len(t_)))
          d_ = lhtml.fromstring(str(t_))  # document

          # lxml cleaner to tidy
          cleaner = Cleaner(style=True)
          c_ = cleaner.clean_html(d_)

          # remove styles, scripts
          # for s_ in d_.xpath("//style"):
          #   s_.getparent().remove(s_)

          # for s_ in d_.xpath("//script"):
          #   s_.getparent().remove(s_)

          # this is the raw text content of the page
          self._html = BeautifulSoup(et.tostring(c_), 'html.parser')
          self._full_text = self._html.get_text("\n")
          logging.info("Length of extracted text: " + str(len(self._full_text)))

          # create list of text components, labeled by type (h, p, etc.)
          # and including the "(U)" paragraph marking
          self._text_list = []
          allowed_elems = ["h1", "h2", "h3", "h4", "h5", "h6", "p"] # "a" not currently supported by python-docx
          for child in self._html.descendants:
            if child.name in allowed_elems:
              # if child.name == 'a':
              #   if "href" in child.attrs:
              #     link = child.attrs["href"]
              #   else:
              #     link = None
              #   if child.get_text() != "" and ( link is not None and link != "#"):
              #     print((child.name, child.get_text(), link))
              # else:
              if len(child.get_text()) > 1:
                if child.name in allowed_elems[:6]:
                  # all header elements converted to level 3 (h3)
                  self._text_list.append(("h", 3, "(U) " + child.get_text()))
                elif child.name == allowed_elems[6]:
                  self._text_list.append(("p", "(U) " + child.get_text()))

  def docxify(self, out_path=WRITE_DIR):
    file_name = re.split("/", self.url)[-1]
    if len(file_name) == 0:
      file_name = re.split("/", self.url)[-2]
    logging.info("Writing to file: " + file_name)

    if self._text_list is None:
      logging.warn("_test_list does not exist!")
    else:
      self.document = Document()
      for text in self._text_list:
        if text[0] == "h":
          self.document.add_heading(text[2], level=text[1])
        elif text[0] == "p":
          self.document.add_paragraph(text[1])

      self.document.save(out_path + file_name + ".docx")

  def pipeline(self, url, html_ids):
    self.fetch(url)
    self.parse(html_ids)
    self.docxify()

##
## IFMAIN
##

if __name__ == '__main__':
  main()

##
## SCRATCH
##

# allowed_elems = ["h1", "h2", "h3", "h4", "h5", "h6", "p"] # "a" not currently supported by python-docx

# doc_components = []

# for child in doc._html.descendants:
#   if child.name in allowed_elems:
#     # if child.name == 'a':
#     #   if "href" in child.attrs:
#     #     link = child.attrs["href"]
#     #   else:
#     #     link = None
#     #   if child.get_text() != "" and ( link is not None and link != "#"):
#     #     print((child.name, child.get_text(), link))
#     # else:
#     if len(child.get_text()) > 0:
#       if child.name in allowed_elems[:6]:
#         # all header elements converted to level 3 (h3)
#         doc_components.append(("h", 3, "(U) " + child.get_text()))
#       elif child.name == allowed_elems[6]:
#         doc_components.append(("p", "(U) " + child.get_text()))

# document = Document()

# for text in doc_components:
#   if text[0] == "h":
#     document.add_heading(text[2], level=text[1])
#   elif text[0] == "p":
#     document.add_paragraph(text[1])

# document.save(file_name)